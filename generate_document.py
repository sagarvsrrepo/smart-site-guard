from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
import os

# Create architecture diagram image
width, height = 1200, 800
img = Image.new('RGB', (width, height), color='white')
draw = ImageDraw.Draw(img)
font = None
try:
    font = ImageFont.truetype('arial.ttf', 18)
except Exception:
    font = ImageFont.load_default()

# Draw boxes and text
boxes = [
    ((100, 120), (340, 220), 'Simulator\nSensors'),
    ((420, 120), (760, 220), 'AWS IoT Core\nRaw Topic'),
    ((840, 120), (1080, 220), 'Fog Processor\nSubscribe Raw'),
    ((420, 320), (760, 420), 'AWS IoT Core\nProcessed Topic'),
    ((420, 520), (760, 620), 'DynamoDB\nSmartSiteGuardEvents'),
    ((840, 520), (1080, 620), 'Flask Dashboard\nRead from DynamoDB'),
]
for top_left, bottom_right, text in boxes:
    draw.rectangle([top_left, bottom_right], outline='black', width=3)
    lines = text.split('\n')
    line_sizes = [draw.textbbox((0, 0), line, font=font) for line in lines]
    line_widths = [bbox[2] - bbox[0] for bbox in line_sizes]
    line_heights = [bbox[3] - bbox[1] for bbox in line_sizes]
    w = max(line_widths)
    h = sum(line_heights) + (len(lines) - 1) * 4
    text_x = top_left[0] + (bottom_right[0] - top_left[0] - w) / 2
    text_y = top_left[1] + (bottom_right[1] - top_left[1] - h) / 2
    current_y = text_y
    for i, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=font)
        line_w = bbox[2] - bbox[0]
        line_h = bbox[3] - bbox[1]
        draw.text((top_left[0] + (bottom_right[0] - top_left[0] - line_w) / 2, current_y), line, fill='black', font=font)
        current_y += line_h + 4

# Draw arrows
arrows = [
    ((340, 170), (420, 170)),
    ((760, 170), (840, 170)),
    ((960, 220), (960, 320)),
    ((640, 420), (640, 520)),
    ((760, 570), (840, 570)),
]
for start, end in arrows:
    draw.line([start, end], fill='black', width=3)
    # arrowhead
    if start[0] != end[0]:
        sign = 1 if end[0] > start[0] else -1
        arrow = [(end[0] - 12 * sign, end[1] - 8), (end[0], end[1]), (end[0] - 12 * sign, end[1] + 8)]
    else:
        sign = 1 if end[1] > start[1] else -1
        arrow = [(end[0] - 8, end[1] - 12 * sign), (end[0], end[1]), (end[0] + 8, end[1] - 12 * sign)]
    draw.polygon(arrow, fill='black')

# Add labels to arrows
label_positions = [
    ((220, 90), 'Publish raw event'),
    ((900, 90), 'Subscribe to raw topic'),
    ((1020, 340), 'Publish processed event'),
    ((700, 470), 'Write event'),
    ((900, 570), 'Load dashboard data'),
]
for pos, label in label_positions:
    draw.text(pos, label, fill='black', font=font)

image_path = 'architecture_diagram.png'
img.save(image_path)

# Create DOCX document
doc = Document()
doc.add_heading('SmartSite Guard Architecture', level=1)
doc.add_paragraph('This document describes the local SmartSite Guard implementation, including the dashboard, simulator, fog processor, and cloud lambda/IoT integration components.')

doc.add_heading('System Overview', level=2)
doc.add_paragraph('The system includes the following components:')
points = [
    'Simulator: generates sensor data and publishes raw events to AWS IoT Core.',
    'AWS IoT Core: routes raw messages on a raw topic and enables fog processor subscriptions.',
    'Fog Processor: subscribes to raw IoT messages, classifies and filters events, stores processed data in DynamoDB, and republishes processed events.',
    'DynamoDB: stores processed sensor events for dashboard visualization and historical queries.',
    'Dashboard: retrieves filtered events from DynamoDB and presents them through charts, summaries, and alerts.',
    'Cloud Lambda: optional API/lambda function for direct ingestion or event querying as a cloud endpoint.',
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Architecture Diagram', level=2)
doc.add_paragraph('The following diagram illustrates the data flow from local sensor simulation through IoT ingestion to fog processing and dashboard storage.')
doc.add_picture(image_path, width=Inches(6))

doc.add_heading('Data Flow', level=2)
flow_steps = [
    ('Sensor Generation', 'The simulator generates events for temperature, gas, noise, vibration, and proximity sensors at regular intervals.'),
    ('Raw Event Publishing', 'The simulator publishes raw sensor events to AWS IoT Core on the configured raw topic.'),
    ('Fog Processing', 'The fog processor subscribes to the raw topic, validates and classifies events, and filters them based on severity.'),
    ('Event Storage', 'Processed events are normalized and stored in DynamoDB for persistent querying by the dashboard.'),
    ('Dashboard Visualization', 'The Flask dashboard reads events from DynamoDB and displays charts, summaries, and recent alerts.'),
    ('Optional Cloud Lambda', 'The cloud lambda can ingest HTTP or event-based payloads and store them in DynamoDB, enabling cloud-based integrations.'),
]
for title, desc in flow_steps:
    doc.add_paragraph(title, style='List Number')
    doc.add_paragraph(desc)

doc.add_heading('Sensor Selection Justification', level=2)
justifications = [
    ('Temperature', 'Detects overheating or abnormal thermal patterns that may indicate equipment failure or fire risk.'),
    ('Gas', 'Monitors air quality and detects hazardous gas buildup, supporting early warning for leaks or toxic conditions.'),
    ('Noise', 'Measures sound levels to detect intrusions, machinery anomalies, or environment disruptions.'),
    ('Vibration', 'Captures mechanical health and structural integrity issues by sensing excessive vibration or instability.'),
    ('Proximity', 'Supports presence detection or object movement in the monitored zone, useful for access control and safety.'),
]
for sensor, reason in justifications:
    doc.add_paragraph(f'{sensor}: {reason}')

doc.add_heading('Local Test Results', level=2)
doc.add_paragraph('The local runtime validation confirmed that all Python modules compile successfully and required imports are available for dashboard, fog, simulator, and cloud lambda components.')

doc.add_paragraph('These checks ensure the project can run locally, assuming AWS credentials and IoT configuration values are provided in the .env file.')

doc.save('SmartSite-Guard-Architecture.docx')
print('Generated SmartSite-Guard-Architecture.docx and architecture_diagram.png')
